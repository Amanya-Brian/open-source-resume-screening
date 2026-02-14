"""Document parser for extracting text from resumes."""

import io
import logging
import re
from pathlib import Path
from typing import Optional

import pdfplumber
import PyPDF2
from docx import Document as DocxDocument

from src.models.schemas import (
    ContactInfo,
    EducationEntry,
    ExperienceEntry,
    ParsedResume,
    ProjectEntry,
)

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for extracting and structuring resume content.

    Supports PDF and DOCX file formats.
    """

    # Common section headers in resumes
    SECTION_PATTERNS = {
        "education": r"(?i)^(education|academic|qualifications?|degrees?)",
        "experience": r"(?i)^(experience|employment|work\s*history|professional\s*experience)",
        "skills": r"(?i)^(skills|technical\s*skills|competencies|technologies)",
        "projects": r"(?i)^(projects|portfolio|personal\s*projects)",
        "certifications": r"(?i)^(certifications?|certificates?|credentials?)",
        "summary": r"(?i)^(summary|objective|profile|about)",
        "languages": r"(?i)^(languages?|language\s*proficiency)",
    }

    # Email pattern
    EMAIL_PATTERN = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"

    # Phone pattern (various formats)
    PHONE_PATTERN = r"[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}"

    # LinkedIn URL pattern
    LINKEDIN_PATTERN = r"(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+"

    # GitHub URL pattern
    GITHUB_PATTERN = r"(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_-]+"

    def __init__(self):
        """Initialize the document parser."""
        self._spacy_model = None

    def parse_file(self, file_path: str) -> ParsedResume:
        """Parse a resume file and extract structured data.

        Args:
            file_path: Path to the resume file

        Returns:
            ParsedResume with extracted data

        Raises:
            ValueError: If file type is not supported
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension == ".pdf":
            text = self.extract_text_from_pdf(file_path)
        elif extension in (".docx", ".doc"):
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        return self.parse_text(text)

    def parse_bytes(self, content: bytes, file_type: str) -> ParsedResume:
        """Parse resume content from bytes.

        Args:
            content: File content as bytes
            file_type: File type (pdf, docx)

        Returns:
            ParsedResume with extracted data
        """
        file_type = file_type.lower().strip(".")

        if file_type == "pdf":
            text = self._extract_pdf_from_bytes(content)
        elif file_type in ("docx", "doc"):
            text = self._extract_docx_from_bytes(content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        return self.parse_text(text)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        text_parts = []

        try:
            # Try pdfplumber first (better formatting)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            if text_parts:
                return "\n\n".join(text_parts)

        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")

        # Fallback to PyPDF2
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise

        return "\n\n".join(text_parts)

    def _extract_pdf_from_bytes(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        text_parts = []

        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            if text_parts:
                return "\n\n".join(text_parts)

        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

        # Fallback
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise

        return "\n\n".join(text_parts)

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n".join(paragraphs)

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_docx_from_bytes(self, content: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            doc = DocxDocument(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n".join(paragraphs)

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def parse_text(self, text: str) -> ParsedResume:
        """Parse extracted text into structured resume data.

        Args:
            text: Raw text content from resume

        Returns:
            ParsedResume with structured data
        """
        # Clean text
        text = self._clean_text(text)

        # Extract sections
        sections = self._extract_sections(text)

        # Build parsed resume
        parsed = ParsedResume(
            raw_text=text,
            contact_info=self._extract_contact_info(text),
            summary=sections.get("summary", ""),
            education=self._extract_education(sections.get("education", "")),
            experience=self._extract_experience(sections.get("experience", "")),
            skills=self._extract_skills(sections.get("skills", text)),
            projects=self._extract_projects(sections.get("projects", "")),
            certifications=self._extract_certifications(sections.get("certifications", "")),
            languages=self._extract_languages(sections.get("languages", "")),
        )

        return parsed

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text.

        Args:
            text: Raw text

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r"[^\w\s.,;:@\-/()'+#]", "", text)
        return text.strip()

    def _extract_sections(self, text: str) -> dict[str, str]:
        """Extract sections from resume text.

        Args:
            text: Resume text

        Returns:
            Dictionary of section name to content
        """
        sections = {}
        lines = text.split("\n")
        current_section = None
        current_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this line is a section header
            section_found = None
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, line):
                    section_found = section_name
                    break

            if section_found:
                # Save previous section
                if current_section:
                    sections[current_section] = "\n".join(current_content)
                current_section = section_found
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_section:
            sections[current_section] = "\n".join(current_content)

        return sections

    def _extract_contact_info(self, text: str) -> ContactInfo:
        """Extract contact information from text.

        Args:
            text: Resume text

        Returns:
            ContactInfo with extracted data
        """
        contact = ContactInfo()

        # Extract email
        email_match = re.search(self.EMAIL_PATTERN, text)
        if email_match:
            contact.email = email_match.group()

        # Extract phone
        phone_match = re.search(self.PHONE_PATTERN, text)
        if phone_match:
            contact.phone = phone_match.group()

        # Extract LinkedIn
        linkedin_match = re.search(self.LINKEDIN_PATTERN, text, re.IGNORECASE)
        if linkedin_match:
            contact.linkedin = linkedin_match.group()

        # Extract GitHub
        github_match = re.search(self.GITHUB_PATTERN, text, re.IGNORECASE)
        if github_match:
            contact.github = github_match.group()

        return contact

    def _extract_education(self, text: str) -> list[EducationEntry]:
        """Extract education entries from text.

        Args:
            text: Education section text

        Returns:
            List of EducationEntry objects
        """
        if not text:
            return []

        entries = []
        lines = text.split("\n")

        # Simple heuristic: look for degree keywords
        degree_patterns = [
            r"(?i)(bachelor|master|phd|doctorate|associate|mba|bs|ba|ms|ma|bsc|msc)",
            r"(?i)(b\.?s\.?|m\.?s\.?|b\.?a\.?|m\.?a\.?|ph\.?d\.?)",
        ]

        current_entry = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if line contains a degree
            has_degree = any(re.search(p, line) for p in degree_patterns)

            if has_degree:
                if current_entry:
                    entries.append(current_entry)

                # Try to parse the line
                current_entry = EducationEntry(
                    institution="",
                    degree=line,
                    field_of_study="",
                )
            elif current_entry:
                # Add to current entry description
                if not current_entry.institution:
                    current_entry.institution = line
                elif not current_entry.field_of_study:
                    current_entry.field_of_study = line

        if current_entry:
            entries.append(current_entry)

        return entries

    def _extract_experience(self, text: str) -> list[ExperienceEntry]:
        """Extract work experience from text.

        Args:
            text: Experience section text

        Returns:
            List of ExperienceEntry objects
        """
        if not text:
            return []

        entries = []
        lines = text.split("\n")

        current_entry = None
        description_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this looks like a job title/company line
            # Usually contains job title keywords or company indicators
            is_new_entry = (
                re.search(r"(?i)(engineer|developer|manager|analyst|designer|intern|lead|senior|junior)", line)
                or re.search(r"\d{4}", line)  # Contains year
            )

            if is_new_entry and len(line) < 200:  # Reasonable title length
                if current_entry:
                    current_entry.description = " ".join(description_lines)
                    entries.append(current_entry)
                    description_lines = []

                current_entry = ExperienceEntry(
                    company="",
                    title=line,
                    description="",
                )
            elif current_entry:
                description_lines.append(line)

        if current_entry:
            current_entry.description = " ".join(description_lines)
            entries.append(current_entry)

        return entries

    def _extract_skills(self, text: str) -> list[str]:
        """Extract skills from text.

        Args:
            text: Skills section or full resume text

        Returns:
            List of skill strings
        """
        if not text:
            return []

        skills = set()

        # Common skill patterns
        skill_keywords = [
            "python", "java", "javascript", "typescript", "c\\+\\+", "c#", "ruby",
            "go", "rust", "swift", "kotlin", "scala", "php", "perl", "r",
            "sql", "nosql", "mongodb", "postgresql", "mysql", "redis", "elasticsearch",
            "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
            "react", "angular", "vue", "node\\.js", "express", "django", "flask", "spring",
            "machine learning", "deep learning", "nlp", "computer vision",
            "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
            "git", "jenkins", "ci/cd", "agile", "scrum", "jira",
            "html", "css", "sass", "webpack", "babel",
            "rest", "graphql", "microservices", "api",
            "linux", "unix", "windows", "macos",
            "communication", "leadership", "teamwork", "problem solving",
        ]

        text_lower = text.lower()

        for skill in skill_keywords:
            if re.search(rf"\b{skill}\b", text_lower):
                # Capitalize properly
                skill_name = skill.replace("\\", "")
                if skill_name == "aws":
                    skill_name = "AWS"
                elif skill_name == "gcp":
                    skill_name = "GCP"
                elif skill_name == "sql":
                    skill_name = "SQL"
                elif skill_name == "nosql":
                    skill_name = "NoSQL"
                elif "." in skill_name:
                    skill_name = skill_name.title()
                else:
                    skill_name = skill_name.capitalize()

                skills.add(skill_name)

        return sorted(list(skills))

    def _extract_projects(self, text: str) -> list[ProjectEntry]:
        """Extract projects from text.

        Args:
            text: Projects section text

        Returns:
            List of ProjectEntry objects
        """
        if not text:
            return []

        projects = []
        lines = text.split("\n")

        current_project = None
        description_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Project titles are usually short and may contain links
            is_project_title = len(line) < 100 and (
                re.search(r"github|gitlab|bitbucket|http", line, re.IGNORECASE)
                or line[0].isupper()
            )

            if is_project_title and not description_lines:
                if current_project:
                    current_project.description = " ".join(description_lines)
                    projects.append(current_project)
                    description_lines = []

                current_project = ProjectEntry(
                    name=line,
                    description="",
                    technologies=[],
                )
            elif current_project:
                description_lines.append(line)

        if current_project:
            current_project.description = " ".join(description_lines)
            projects.append(current_project)

        return projects

    def _extract_certifications(self, text: str) -> list[str]:
        """Extract certifications from text.

        Args:
            text: Certifications section text

        Returns:
            List of certification names
        """
        if not text:
            return []

        certifications = []
        lines = text.split("\n")

        for line in lines:
            line = line.strip()
            if line and len(line) < 200:
                certifications.append(line)

        return certifications

    def _extract_languages(self, text: str) -> list[str]:
        """Extract languages from text.

        Args:
            text: Languages section text

        Returns:
            List of language names
        """
        if not text:
            return []

        languages = []

        # Common language patterns
        language_keywords = [
            "english", "french", "spanish", "german", "chinese", "mandarin",
            "japanese", "korean", "arabic", "hindi", "portuguese", "russian",
            "italian", "dutch", "swahili", "kinyarwanda",
        ]

        text_lower = text.lower()

        for lang in language_keywords:
            if lang in text_lower:
                languages.append(lang.capitalize())

        return languages
